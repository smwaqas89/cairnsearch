"""DOCX text extraction."""
from pathlib import Path
from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseExtractor, ExtractionResult


class DocxExtractor(BaseExtractor):
    """Extract text from DOCX files."""

    @property
    def supported_extensions(self) -> list[str]:
        # Only handle .docx - old .doc files need different handling
        return [".docx"]

    def extract(self, file_path: Path) -> ExtractionResult:
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Extract tables
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Combine all text
            all_text = paragraphs + table_texts
            text = "\n\n".join(all_text)

            # Extract metadata from core properties
            core_props = doc.core_properties
            
            return ExtractionResult(
                success=True,
                text=text,
                metadata={
                    "title": core_props.title,
                    "author": core_props.author,
                    "created_date": core_props.created.isoformat() if core_props.created else None,
                    "modified_date": core_props.modified.isoformat() if core_props.modified else None,
                },
                extraction_method="direct"
            )
        except PackageNotFoundError:
            return ExtractionResult(
                success=False, 
                error="Invalid or corrupted DOCX file"
            )
        except Exception as e:
            return ExtractionResult(success=False, error=str(e))


class DocExtractor(BaseExtractor):
    """Extract text from old .doc files using antiword or textract."""

    @property
    def supported_extensions(self) -> list[str]:
        return [".doc"]

    def extract(self, file_path: Path) -> ExtractionResult:
        """Try multiple methods to extract from old .doc format."""
        
        # Method 1: Try using antiword (if installed)
        text = self._try_antiword(file_path)
        if text:
            return ExtractionResult(
                success=True,
                text=text,
                extraction_method="antiword"
            )
        
        # Method 2: Try using catdoc (if installed)
        text = self._try_catdoc(file_path)
        if text:
            return ExtractionResult(
                success=True,
                text=text,
                extraction_method="catdoc"
            )
        
        # Method 3: Try reading as text (some .doc files are actually RTF or text)
        text = self._try_raw_text(file_path)
        if text:
            return ExtractionResult(
                success=True,
                text=text,
                extraction_method="raw_text"
            )
        
        return ExtractionResult(
            success=False,
            error="Cannot extract .doc file. Install antiword: brew install antiword"
        )
    
    def _try_antiword(self, file_path: Path) -> str | None:
        """Try extracting with antiword."""
        try:
            import subprocess
            result = subprocess.run(
                ["antiword", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None
    
    def _try_catdoc(self, file_path: Path) -> str | None:
        """Try extracting with catdoc."""
        try:
            import subprocess
            result = subprocess.run(
                ["catdoc", str(file_path)],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except (FileNotFoundError, subprocess.TimeoutExpired):
            pass
        return None
    
    def _try_raw_text(self, file_path: Path) -> str | None:
        """Try reading raw text from file."""
        try:
            content = file_path.read_bytes()
            # Check if it might be RTF
            if content.startswith(b'{\\rtf'):
                # Strip RTF tags (basic)
                import re
                text = content.decode('utf-8', errors='ignore')
                text = re.sub(r'\\[a-z]+\d* ?', '', text)
                text = re.sub(r'[{}]', '', text)
                text = text.strip()
                if len(text) > 50:
                    return text
        except Exception:
            pass
        return None
